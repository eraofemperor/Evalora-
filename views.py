from django.shortcuts import render, redirect
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth import login, logout, authenticate
from .forms import CreateUserForm
from django.contrib import messages
from .forms import LoginForm
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib.auth.hashers import make_password
from .forms import PasswordResetForm
from .forms import AnswerSheetUploadForm
from .forms import EvaluationForm
from .models import Subject,AnswerSheet, Evaluation, EvaluationDetail
from django.shortcuts import render, get_object_or_404
from django.http import HttpResponse 
from django.http import JsonResponse
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from django.shortcuts import render
from django.http import HttpResponse
from django.db.models import Avg
import io
import csv
import xlsxwriter
from docx import Document
from reportlab.pdfgen import canvas


def register_view(request):
    if request.method == 'POST':
        form = CreateUserForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Account created successfully! Please login.")
            return redirect('login_view')
        else:
            print(form.errors)
    else:
        form = CreateUserForm()
    return render(request, 'register.html', {'form': form})

def login_view(request):
    if request.method == "POST":
        form = LoginForm(request, data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get("username")
            password = form.cleaned_data.get("password")
            user = authenticate(username=username, password=password)
            if user is not None:
                login(request, user)
                messages.success(request, f"Welcome {username}! You are now logged in.")
                return redirect("index") 
            else:
                messages.error(request, "Invalid username or password.")
        else:
            messages.error(request, "Invalid username or password.")
    else:
        form = LoginForm()
    return render(request, "login.html", {"form": form})


def logout_view(request):
    logout(request)
    messages.info(request, "You have logged out successfully.")
    return redirect("login")



@login_required(login_url='login')  # Redirects to login if not logged in
def index(request):
    return render(request, 'index.html')

def reset_password(request):
    if request.method == "POST":
        form = PasswordResetForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data['username']
            new_password = form.cleaned_data['password1']  # use password1
            try:
                user = User.objects.get(username=username)
                user.set_password(new_password)
                user.save()
                messages.success(request, "Password reset successful. You can now log in.")
                return redirect('login_view')
            except User.DoesNotExist:
                messages.error(request, "Username does not exist.")
    else:
        form = PasswordResetForm()
    return render(request, 'reset_password.html', {'form': form})

@login_required(login_url='login')
def upload_view(request):
    if request.method == "POST":
        rollno = request.POST.get("rollno")
        subject_name = request.POST.get("subject_name")
        uploaded_file = request.FILES.get("answer_sheet")  # keep the form field name

        if not rollno or not subject_name or not uploaded_file:
            messages.error(request, "All fields are required.")
            return redirect("upload_view")

        # Create subject if it doesn’t exist
        subject, created = Subject.objects.get_or_create(name=subject_name)

        # Save the answer sheet (field name = file)
        AnswerSheet.objects.create(
            rollno=rollno,
            subject=subject,
            file=uploaded_file
        )

        messages.success(request, f"Answer sheet uploaded for {subject.name}")
        return redirect("upload_view")

    subjects = Subject.objects.all()
    return render(request, "upload.html", {"subjects": subjects})

@login_required(login_url='login')
def evaluate(request):
    if request.method == "POST":
        form = EvaluationForm(request.POST)
        # --- Get AnswerSheet ---
        answer_sheet_id = request.POST.get("answer_sheet_id")
        if not answer_sheet_id or answer_sheet_id.strip() == "":
            messages.error(request, "Please select an answer sheet.")
            return redirect("evaluate")
        
        try:
            answer_sheet = AnswerSheet.objects.get(id=answer_sheet_id)
        except AnswerSheet.DoesNotExist:
            messages.error(request, "Answer sheet not found.")
            return redirect("evaluate")
        
        # --- Create Evaluation ---
        evaluation = Evaluation.objects.create(
            answer_sheet=answer_sheet,
            evaluator=request.user,
            total_marks=0
        )
        total_marks = 0

        # --- Q1–Q10 ---
        for i in range(1, 11):
            mark = request.POST.get(f"mark_q4_{i}") or 0
            comment = request.POST.get(f"comment_q4_{i}") or ""
            EvaluationDetail.objects.create(
                evaluation=evaluation,
                question_number=i,
                mark=int(mark),
                comment=comment
            )
            total_marks += int(mark)

        # --- Q11–Q22 ---
        selected_questions = request.POST.getlist("q10_select")
        if len(selected_questions) != 6:
            evaluation.delete()
            messages.error(request, "Select exactly 6 questions from Q11–Q22.")
            return redirect("evaluate")

        for q_str in selected_questions:
            q = int(q_str)
            mark = request.POST.get(f"mark_q10_{q}") or 0
            comment = request.POST.get(f"comment_q10_{q}") or ""
            EvaluationDetail.objects.create(
                evaluation=evaluation,
                question_number=q,
                mark=int(mark),
                comment=comment
            )
            total_marks += int(mark)

        # --- Total marks validation ---
        if total_marks > 100:
            evaluation.delete()
            messages.error(request, "Total marks cannot exceed 100.")
            return redirect("evaluate")

        evaluation.total_marks = total_marks
        evaluation.save()
        messages.success(request, "Evaluation submitted successfully!")
        return redirect("evaluate")

    else:
        form = EvaluationForm()

    context = {
        "form": form,
        "answer_sheets": AnswerSheet.objects.all(),
        "questions_4": range(1, 11),
        "questions_10": range(11, 23)
    }
    return render(request, "evaluation.html", context)

@login_required(login_url='login')
def ocr_convert(request, sheet_id):
    sheet = get_object_or_404(AnswerSheet, id=sheet_id)
    # TODO: integrate ML OCR model
    return JsonResponse({"result": f"OCR processing for {sheet.id} will happen here."})

# ----------------- Download Page -----------------
@login_required(login_url='login')
@login_required(login_url='login')
def download(request):
    subjects = Subject.objects.all()

    # Insights
    evaluations = Evaluation.objects.all().select_related("answer_sheet__student")
    topper = evaluations.order_by("-total_marks").first()
    least = evaluations.order_by("total_marks").first()
    average_marks = evaluations.aggregate(avg=Avg("total_marks"))["avg"] or 0
    total = evaluations.count()
    passed = evaluations.filter(total_marks__gte=40).count()  # pass mark assumption
    pass_percentage = (passed / total * 100) if total > 0 else 0

    # ✅ Safe handling for chart labels
    chart_labels = []
    chart_data = []
    for e in evaluations:
        if e.answer_sheet and e.answer_sheet.student:
            chart_labels.append(e.answer_sheet.student.name)
            chart_data.append(e.total_marks)
        else:
            chart_labels.append("Unknown")
            chart_data.append(e.total_marks if e.total_marks else 0)

    context = {
        "subjects": subjects,
        "topper": topper,
        "least": least,
        "average_marks": round(average_marks, 2),
        "pass_percentage": round(pass_percentage, 2),
        "chart_labels": chart_labels,
        "chart_data": chart_data,
    }
    return render(request, "download.html", context)


# ----------------- ALL Results -----------------
@login_required(login_url='login')
def download_all_results_pdf(request):
    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = 'attachment; filename="all_results.pdf"'

    p = canvas.Canvas(response)
    p.setFont("Helvetica", 12)
    p.drawString(100, 800, "All Student Results")

    y = 760
    for eval in Evaluation.objects.select_related("answer_sheet__student"):
        line = f"{eval.answer_sheet.student.name} - {eval.total_marks}"
        p.drawString(100, y, line)
        y -= 20

    p.showPage()
    p.save()
    return response

@login_required(login_url='login')
def download_all_results_excel(request):
    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output)
    worksheet = workbook.add_worksheet("Results")

    worksheet.write(0, 0, "Student")
    worksheet.write(0, 1, "Marks")

    row = 1
    for eval in Evaluation.objects.select_related("answer_sheet__student"):
        worksheet.write(row, 0, eval.answer_sheet.student.name)
        worksheet.write(row, 1, eval.total_marks)
        row += 1

    workbook.close()
    output.seek(0)

    response = HttpResponse(output.read(), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response["Content-Disposition"] = 'attachment; filename="all_results.xlsx"'
    return response

@login_required(login_url='login')
def download_all_results_word(request):
    document = Document()
    document.add_heading("All Student Results", 0)

    for eval in Evaluation.objects.select_related("answer_sheet__student"):
        document.add_paragraph(f"{eval.answer_sheet.student.name}: {eval.total_marks}")

    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = 'attachment; filename="all_results.docx"'
    document.save(response)
    return response


# ----------------- SUBJECT-wise -----------------
@login_required(login_url='login')
def download_subject_results_pdf(request, subject_id):
    subject = get_object_or_404(Subject, id=subject_id)
    evaluations = Evaluation.objects.filter(answer_sheet__subject=subject).select_related("answer_sheet__student")

    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{subject.name}_results.pdf"'

    p = canvas.Canvas(response)
    p.setFont("Helvetica", 12)
    p.drawString(100, 800, f"Results for {subject.name}")

    y = 760
    for eval in evaluations:
        line = f"{eval.answer_sheet.student.name} - {eval.total_marks}"
        p.drawString(100, y, line)
        y -= 20

    p.showPage()
    p.save()
    return response

@login_required(login_url='login')
def download_subject_results_excel(request, subject_id):
    subject = get_object_or_404(Subject, id=subject_id)
    evaluations = Evaluation.objects.filter(answer_sheet__subject=subject).select_related("answer_sheet__student")

    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output)
    worksheet = workbook.add_worksheet(subject.name)

    worksheet.write(0, 0, "Student")
    worksheet.write(0, 1, "Marks")

    row = 1
    for eval in evaluations:
        worksheet.write(row, 0, eval.answer_sheet.student.name)
        worksheet.write(row, 1, eval.total_marks)
        row += 1

    workbook.close()
    output.seek(0)

    response = HttpResponse(output.read(), content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response["Content-Disposition"] = f'attachment; filename="{subject.name}_results.xlsx"'
    return response


# ----------------- STUDENT -----------------
@login_required(login_url='login')
def search_student_result(request):
    rollno = request.GET.get("rollno")
    if not rollno:
        return render(request, "student_result.html", {"error": "Please enter roll number"})

    evaluation = Evaluation.objects.filter(answer_sheet__student__student_id=rollno).first()
    if not evaluation:
        return render(request, "download.html", {"error": "No result found for this roll number"})

    context = {
        "student": evaluation.answer_sheet.student,
        "marks": evaluation.total_marks,
        "subject": evaluation.answer_sheet.subject,
    }
    return render(request, "download.html", context)

